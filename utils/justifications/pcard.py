#!/bin/python3
import docx, pathlib, os, random, pytz, json, shutil
from datetime import datetime
from main import path_to_members, path_to_projects
from ..utils.helper import drop_a_line

#####

with open(path_to_members) as incoming:
    members = json.load(incoming)

with open(path_to_projects) as incoming:
    projects = json.load(incoming)

# Current time in Pacific
now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m/%d/%Y")
right_now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m_%d_%Y")


##########


class PCard:
    def __init__(
        self, here, charge_to_card, j_short, j_long, j_why, who, when, project
    ):

        self.timestamp = right_now

        self.base_path = os.path.join(here, "files/justifications")
        self.input_path = os.path.join(here, "files/templates")
        self.output_path = os.path.join(self.base_path, right_now)

        pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

        self.output_filename = f"pcard_{charge_to_card}_zaki.docx"

        self.j_short = j_short
        self.j_long = j_long
        self.j_why = j_why
        self.who = members[who]
        self.when = when
        self.project = projects[project]

        self._charge = charge_to_card

    def load_template(self):
        return docx.Document(os.path.join(self.input_path, "P-Card.docx"))

    def write_justification(self):
        template = self.load_template()

        template.paragraphs[3].text = template.paragraphs[3].text.format(
            self.j_short,
            self.project["award"],
            self.who["full_name"],
            self.who["title"],
            self.j_long,
            self.when,
            self.j_why,
            self.project["funding-string"],
        )

        header = template.sections[0].header
        head = header.paragraphs[0]
        head.text = f"\n\nCreated {now}"

        template.save(os.path.join(self.output_path, self.output_filename))

        drop_a_line(self.output_path)

        self.gunzip()

    def gunzip(self):
        out_path = os.path.join(
            self.base_path, f"SSNL-Justification-{right_now}-{self._charge}"
        )

        shutil.make_archive(out_path, "zip", self.output_path)
