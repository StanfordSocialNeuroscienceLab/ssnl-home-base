#!/bin/python3
import docx, pathlib, os, pytz, json, shutil
from datetime import datetime
from main import path_to_members, path_to_projects

from ..base.helper import drop_a_line

#####

with open(path_to_members) as incoming:
    members = json.load(incoming)

with open(path_to_projects) as incoming:
    projects = json.load(incoming)

# Current time in Pacific
now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m/%d/%Y")
right_now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m_%d_%Y")


##########


class Reocurring:
    def __init__(self, here, charge, date_of_charge):
        self.charge = charge
        self.date = date_of_charge

        today_f = datetime.now(pytz.timezone("US/Pacific")).strftime(
            "%b_%d_%Y_%H_%M_%S"
        )

        self.base_path = os.path.join(here, "files/justifications")
        self.template_path = os.path.join(here, "files/templates/reoccuring")
        self.output_path = os.path.join(self.base_path, today_f)

        if not os.path.exists(self.output_path):
            pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

        self.set_conventions()

    def set_conventions(self):
        if self.charge == "AWS_SCP":
            filename = "pcard_AWS-SCP_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_AWS.docx")

        elif self.charge == "AWS_JM":
            filename = "pcard_AWS-JM_zaki.docx"
            filepath = os.path.join(self.template_path, "JM_AWS.docx")

        elif self.charge == "Mailchimp":
            filename = "pcard_87.00_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_Mailchimp.docx")

        elif self.charge == "Forge":
            filename = "pcard_12.00_zaki.docx"
            filepath = os.path.join(self.template_path, "JM_Forge.docx")

        elif self.charge == "Buzzsprout":
            filename = "pcard_27.00_zaki.docx"
            filepath = os.path.join(self.template_path, "Buzzsprout.docx")

        elif self.charge == "PythonAnywhere":
            filename = "pcard_10.00_zaki.docx"
            filepath = os.path.join(self.template_path, "PythonAnywhere.docx")

        elif self.charge == "SCP_Adobe":
            filename = "pcard_19.99_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_Adobe.docx")

        self.filename = filename
        self.filepath = filepath

    def write_justification(self):

        doc = docx.Document(self.filepath)
        doc.paragraphs[3].text = doc.paragraphs[3].text.format(self.date)

        header = doc.sections[0].header
        head = header.paragraphs[0]
        head.text = "\n\nCreated {}".format(now)

        doc.save(os.path.join(self.output_path, self.filename))

        drop_a_line(self.output_path)

        self.gunzip()

    def gunzip(self):

        out_path = os.path.join(self.base_path, f"SSNL-{self.charge}-{right_now}")

        shutil.make_archive(out_path, "zip", self.output_path)
