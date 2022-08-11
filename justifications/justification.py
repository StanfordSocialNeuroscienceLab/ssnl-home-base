#!/bin/python3

# --- Imports
import docx, pathlib, os, random, pytz, json, shutil
from datetime import datetime

from main import path_to_members, path_to_projects

with open(path_to_members) as incoming:
      members = json.load(incoming)

with open(path_to_projects) as incoming:
      projects = json.load(incoming)

# Current time in Pacific
now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m/%d/%Y")
right_now = datetime.now(pytz.timezone("US/Pacific")).strftime("%m_%d_%Y")


# --- Have a nice day!
def drop_a_line(path):

      options = ["Have a nice day!",
                 "Keep up the great work!",
                 "You're doing great!",
                 "You are the sun and the moon!",
                 "You are absolute magic!",
                 "Thank you for being you!",
                 "Make today great!",
                 "Make it a great day!",
                 "Do something nice for yourself today!",
                 "You are a light everywhere you go!",
                 "You can accomplish anything you set your mind to!",
                 "Everyone needs a friend like you!"]

      with open(os.path.join(path, "README.txt"), "w") as file:
            message = random.choice(options)

            file.write("PLEASE READ\n\n")
            file.write(message)



# --- Object definitions
class PCard:
      def __init__(self, here, charge_to_card, j_short, j_long, j_why, who, when, project):

            today_f = datetime.now().strftime("%b_%d_%Y_%H_%M_%S")

            self.base_path = os.path.join(here, "files/justifications")

            self.input_path = os.path.join(here, "files/templates")
            self.output_path = os.path.join(self.base_path, today_f)

            if not os.path.exists(self.output_path):
                  pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

            self.output_filename = f"pcard_{charge_to_card}_zaki.docx"

            self.j_short = j_short
            self.j_long = j_long
            self.j_why = j_why
            self.who = members[who]
            self.when = when
            self.project = projects[project]

            self._charge = charge_to_card


      def load_template(self):
            return docx.Document(os.path.join(self.input_path, "P-Card.docx"))


      def write_justification(self):
            template = self.load_template()

            template.paragraphs[3].text = template.paragraphs[3].text.format(self.j_short,
                                                                             self.project["award"],
                                                                             self.who["full_name"],
                                                                             self.who["title"],
                                                                             self.j_long,
                                                                             self.when,
                                                                             self.j_why,
                                                                             self.project["funding-string"])


            header = template.sections[0].header
            head = header.paragraphs[0]
            head.text = f"\n\nCreated {now}"

            template.save(os.path.join(self.output_path, self.output_filename))

            drop_a_line(self.output_path)

            self.gunzip()


      def gunzip(self):
            out_path = os.path.join(self.base_path, f"SSNL-Justification-{right_now}-{self._charge}")

            shutil.make_archive(out_path,
                                "zip",
                                self.output_path)



class Reimbursement:
      def __init__(self, here, charge_to_card, j_short, j_long, j_why, who, when, project):
            today_f = datetime.now().strftime("%b_%d_%Y_%H_%M_%S")

            self.base_path = os.path.join(here, "files/justifications")

            self.input_path = os.path.join(here, "files/templates")
            self.output_path = os.path.join(self.base_path, today_f)

            if not os.path.exists(self.output_path):
                  pathlib.Path(self.output_path).mkdir(
                        exist_ok=True, parents=True)

            self.output_filename = f"reimbursement_{charge_to_card}_zaki.docx"

            self.j_short = j_short
            self.j_long = j_long
            self.j_why = j_why
            self.who = members[who]
            self.when = when
            self.project = projects[project]

            self._charge = charge_to_card


      def load_template(self):
            return docx.Document(os.path.join(self.input_path, "Reimbursement.docx"))


      def write_justification(self):
            template = self.load_template()

            template.paragraphs[3].text = template.paragraphs[3].text.format(self.j_short,
                                                                             self.project["award"],
                                                                             self.who["full_name"],
                                                                             self.who["title"],
                                                                             self.who["employee_number"],
                                                                             self.j_long,
                                                                             self.when,
                                                                             self.j_why,
                                                                             self.project["funding-string"])


            header = template.sections[0].header
            head = header.paragraphs[0]
            head.text = f"\n\nCreated {now}"

            template.save(os.path.join(self.output_path, self.output_filename))

            drop_a_line(self.output_path)

            self.gunzip()


      def gunzip(self):
            out_path = os.path.join(self.base_path, 
                                    f"SSNL-Reimbursement-{right_now}-{self._charge}")

            shutil.make_archive(out_path,
                                "zip",
                                self.output_path)



class Reocurring:
      def __init__(self, here, charge, date_of_charge):
            self.charge = charge
            self.date = date_of_charge

            today_f = datetime.now().strftime("%b_%d_%Y_%H_%M_%S")

            self.base_path = os.path.join(here, "files/justifications")
            self.template_path = os.path.join(here, "files/templates/reoccuring")
            self.output_path = os.path.join(self.base_path, today_f)

            if not os.path.exists(self.output_path):
                  pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

            self.set_conventions()


      def set_conventions(self):
            if self.charge == "AWS_SCP":
                  filename = "pcard_AWS-SCP_zaki.docx"
                  filepath = os.path.join(self.template_path, "SCP_AWS.docx")

            elif self.charge == "AWS_JM":
                  filename = "pcard_AWS-JM_zaki.docx"
                  filepath = os.path.join(self.template_path, "JM_AWS.docx")

            elif self.charge == "Mailchimp":
                  filename = "pcard_87.00_zaki.docx"
                  filepath = os.path.join(self.template_path, "SCP_Mailchimp.docx")

            elif self.charge == "Forge":
                  filename = "pcard_12.00_zaki.docx"
                  filepath = os.path.join(self.template_path, "JM_Forge.docx")

            elif self.charge == "Buzzsprout":
                  filename = "pcard_27.00_zaki.docx"
                  filepath = os.path.join(self.template_path, "Buzzsprout.docx")

            elif self.charge == "PythonAnywhere":
                  filename = "pcard_10.00_zaki.docx"
                  filepath = os.path.join(self.template_path, "PythonAnywhere.docx")

            elif self.charge == "SCP_Adobe":
                  filename = "pcard_19.99_zaki.docx"
                  filepath = os.path.join(self.template_path, "SCP_Adobe.docx")

            
            self.filename = filename
            self.filepath = filepath


      def write_justification(self):

            doc = docx.Document(self.filepath)
            doc.paragraphs[3].text = doc.paragraphs[3].text.format(self.date)

            header = doc.sections[0].header
            head = header.paragraphs[0]
            head.text = "\n\nCreated {}".format(now)

            doc.save(os.path.join(self.output_path, self.filename))

            drop_a_line(self.output_path)

            self.gunzip()


      def gunzip(self):
            out_path = os.path.join(self.base_path, 
                                    f"SSNL-{self.charge}-{right_now}")

            shutil.make_archive(out_path, 
                                "zip", 
                                self.output_path) 