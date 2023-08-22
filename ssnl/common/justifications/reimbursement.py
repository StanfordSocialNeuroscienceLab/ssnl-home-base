#!/bin/python3
import docx
import pathlib
import os
import pytz
import shutil
from datetime import datetime

from ssnl.common.utils import drop_a_line
from .base import FinanceObject

##########


class Reimbursement(FinanceObject):
    def __init__(
        self, here, charge_to_card, j_short, j_long, j_why, who, when, project, where
    ):
        super().__init__()

        today_f = datetime.now(pytz.timezone("US/Pacific")).strftime(
            "%b_%d_%Y_%H_%M_%S"
        )

        self.base_path = os.path.join(here, "files/justifications")

        self.input_path = os.path.join(here, "files/templates")
        self.output_path = os.path.join(self.base_path, today_f)

        if not os.path.exists(self.output_path):
            pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

        self.output_filename = f"reimbursement_{charge_to_card}_zaki.docx"

        self.j_short = j_short
        self.j_long = j_long
        self.j_why = j_why
        self.who = self.member_dictionary[who]

        # If user passes in "TODAY" we won't overwrite the attribute
        if not when.strip().upper() == "TODAY":
            self.when = when

        self.where = where
        self.project = self.project_dictioanry[project]

        self._charge = charge_to_card

        self.timestamp = self.__timestamp

    def load_template(self):
        return docx.Document(os.path.join(self.input_path, "Reimbursement.docx"))

    def write_justification(self):
        template = self.load_template()

        template.paragraphs[3].text = template.paragraphs[3].text.format(
            self.j_short,
            self.project["award"],
            self.who["full_name"],
            self.who["title"],
            self.who["employee_number"],
            self.j_long,
            self.when,
            self.where,
            self.j_why,
            self.project["funding-string"],
        )

        template.paragraphs[3].text = (
            template.paragraphs[3].text.replace('"', "").replace("'", "")
        )

        header = template.sections[0].header
        head = header.paragraphs[0]
        head.text = f"\n\nCreated {self.timestamp.strftime('%m/%d/%Y')}"

        template.save(os.path.join(self.output_path, self.output_filename))

        drop_a_line(self.output_path)

        self.gunzip()

    def gunzip(self):
        out_path = os.path.join(
            self.base_path,
            f"SSNL-Reimbursement-{self.timestamp.strftime('%m_%d_%Y')}-{self._charge}",
        )

        shutil.make_archive(out_path, "zip", self.output_path)
