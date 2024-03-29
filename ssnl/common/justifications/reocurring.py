#!/bin/python3
import docx
import pathlib
import os
import pytz
import shutil
from datetime import datetime

from ssnl.common.utils import drop_a_line
from .base import FinanceObject

##########


class Reocurring(FinanceObject):
    def __init__(self, here, charge, date_of_charge):
        super().__init__()

        self.charge = charge

        if not date_of_charge.strip().upper() == "TODAY":
            self.when = date_of_charge

        self.timestamp = datetime.now(pytz.timezone("US/Pacific"))

        self.base_path = os.path.join(here, "files/justifications")
        self.template_path = os.path.join(here, "files/templates/reoccuring")
        self.output_path = os.path.join(
            self.base_path, self.timestamp.strftime("%b_%d_%Y_%H_%M_%S")
        )

        self.output_name = f"SSNL-{self.charge}-{self.timestamp.strftime('%m_%d_%Y')}"

        if not os.path.exists(self.output_path):
            pathlib.Path(self.output_path).mkdir(exist_ok=True, parents=True)

        self.set_conventions()

    def set_conventions(self):
        if self.charge == "AWS_SCP":
            filename = "pcard_AWS-SCP_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_AWS.docx")

        elif self.charge == "AWS_JM":
            filename = "pcard_AWS-JM_zaki.docx"
            filepath = os.path.join(self.template_path, "JM_AWS.docx")

        elif self.charge == "Mailchimp":
            filename = "pcard_87.00_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_Mailchimp.docx")

        elif self.charge == "Forge":
            filename = "pcard_12.00_zaki.docx"
            filepath = os.path.join(self.template_path, "JM_Forge.docx")

        elif self.charge == "Buzzsprout":
            filename = "pcard_27.00_zaki.docx"
            filepath = os.path.join(self.template_path, "Buzzsprout.docx")

        elif self.charge == "PythonAnywhere":
            filename = "pcard_10.00_zaki.docx"
            filepath = os.path.join(self.template_path, "PythonAnywhere.docx")

        elif self.charge == "SCP_Adobe":
            filename = "pcard_19.99_zaki.docx"
            filepath = os.path.join(self.template_path, "SCP_Adobe.docx")

        self.filename = filename
        self.filepath = filepath

    def write_justification(self):
        doc = docx.Document(self.filepath)

        # Format date
        doc.paragraphs[3].text = doc.paragraphs[3].text.format(self.when)

        # Strip out weird characters for Tracy
        doc.paragraphs[3].text = (
            doc.paragraphs[3].text.replace('"', "").replace("'", "")
        )

        header = doc.sections[0].header
        head = header.paragraphs[0]
        head.text = "\n\nCreated {}".format(self.timestamp.strftime("%m/%d/%Y"))

        doc.save(os.path.join(self.output_path, self.filename))

        drop_a_line(self.output_path)

        self.gunzip()

    def gunzip(self):
        out_path = os.path.join(self.base_path, self.output_name)

        shutil.make_archive(out_path, "zip", self.output_path)
